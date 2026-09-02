#!/usr/bin/env python3
"""Pour camera_ready.md into the real DVCon Europe 2026 template.

Unlike build_docx.py (which recreates the look with hand-defined styles), this
opens the official template .docx and uses it as the base document, so the
result inherits, verbatim:

  - the DVCon Europe 2026 logo in the page header (word/media/image1.png)
  - the page-number footer
  - A4 with the template's own margins (L/R 1.00in, top 1.44in, bottom 1.00in)
  - every named style: 'paper title', 'paper subtitle', 'Author', 'Abstract',
    'key words', 'Heading 1/2/5', 'figure caption', 'references', 'table head',
    'table col head', 'table copy', 'bullet list'

IMPORTANT -- the template AUTO-NUMBERS several styles through numbering.xml:
    Heading 1     numId 4   upperRoman  '%1.'      -> "I.", "II." ...
    Heading 2     numId 5   upperLetter '%2.'      -> "A.", "B." ...
    table head    numId 9   upperRoman  'TABLE %1. '
    references    numId 8   decimal     '[%1]'
So the manual numbering in the markdown ("## I. Introduction", "### A. ...",
"**Table I - ...**", "[1] ...") MUST be stripped, or it renders doubled as
"I. I. INTRODUCTION". Figure captions are NOT auto-numbered, so "Figure 1."
is written out.

Usage: build_dvcon_docx.py camera_ready.md out.docx [template.docx]
"""
import re
import sys
import tempfile
from pathlib import Path

from PIL import Image, ImageChops

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SRC = Path(sys.argv[1])
OUT = Path(sys.argv[2])
TPL = Path(sys.argv[3]) if len(sys.argv) > 3 else Path(
    "/home/jk/Downloads/2026_projects/"
    "dvcon_europe_Engineer_Paper_Template_2026_75c0203916 (4).docx")
BASE = SRC.parent

doc = Document(str(TPL))
sec = doc.sections[0]
TEXT_W = (sec.page_width - sec.left_margin - sec.right_margin) / 914400   # EMU -> in

# Figure widths as a fraction of the text measure. The template's measure is
# 6.27in, and all three figures are wide (aspect 0.39-0.54), so even at full
# width none exceeds 3.4in tall -- there is no reason to shrink them, and at
# the old 0.66 default the labels inside Figure 2's block diagram were too
# small to read in print.
#
# Camera-ready is capped at 8 pages, so the figures are scaled back from the
# full measure they used at draft stage. figure1 is a wide two-panel flow
# comparison (aspect 0.39), so 0.88 costs it almost no legibility while
# returning roughly a third of a page across the document.
FIG_FRAC = {
    "figure1_flow_comparison.png": 0.66,
    "figure2_flows_combined.png": 0.78,
    "figure4_coverage_closure.png": 0.88,
}
FIG_FRAC_DEFAULT = 0.88

# Relative column weights per table; scaled to TEXT_W at build time so they
# always fit the template's narrower measure. One entry per table, and the
# entry's length MUST equal that table's column count -- see set_widths.
TABLE_WEIGHTS = {
    # "Tooling cost" holds short values (Free, Engineer time, ...), so its
    # measure goes to the wrapping "closes the gap?" column.
    "I":   [1.20, 1.10, 0.70, 2.00],
    # Table II's four data columns only ever hold "156 / 0"-shaped values, so
    # they were far wider than needed while the feature column wrapped to two
    # lines on 8 of 17 rows. Narrowing the data columns and giving the measure
    # to the scenario and feature columns un-wraps those rows -- worth ~80pt
    # against the 8-page cap, with no cell content changed.
    #
    # Column 0 must stay >= ~1.6in: the scenario names are monospaced code
    # spans, and the longest ('sw_breakpoint_progbuf', 21 chars) breaks
    # mid-identifier if the column is narrower, which reads as a typo.
    "II":  [1.62, 1.85, 0.70, 0.68, 0.70, 0.68],
    "III": [0.85, 1.15, 1.15, 1.25, 1.35, 0.95],
    "IV":  [1.55, 1.00],
    # Same rebalance as Table II: "Appeared on" only ever holds Both/CVA6
    # only/Ibex only, so its measure goes to the wrapping prose column.
    "V":   [1.35, 0.50, 1.05, 2.10],
    "VI":  [1.55, 0.85, 0.85, 0.62, 0.90, 0.90, 0.62],
}

# Zero-based indices of the body columns to centre. Data columns (counts,
# timings, versions) read as left-hugging in a wide column; prose columns must
# stay left-aligned, since centred running text is unreadable. Header cells are
# centred already by the template's 'table col head' style.
TABLE_CENTRE_COLS = {
    "II":  {2, 3, 4, 5},
    "III": {1, 2, 3, 4, 5},
    "V":   {1},
    "VI":  {1, 2, 3, 4, 5, 6},
}


# --------------------------------------------------------------------------
# Strip the template's own body.
#
# Careful: the template has three sectPr elements. The one carrying the
# headerReference/footerReference (the DVCon logo and the page-number footer)
# is NOT the body-level one -- it is nested in a paragraph's pPr, so clearing
# the body throws the logo and page numbers away. Capture those references
# first, then graft them onto the surviving body-level sectPr.
# --------------------------------------------------------------------------
body = doc.element.body
HDR_FTR = (qn("w:headerReference"), qn("w:footerReference"))

refs = []
for sp in body.iter(qn("w:sectPr")):
    found = [c for c in sp if c.tag in HDR_FTR]
    if found:
        refs = found
        break

for child in list(body):
    if not child.tag.endswith("}sectPr"):
        body.remove(child)

final_sect = body.find(qn("w:sectPr"))
if final_sect is not None and refs:
    have = {c.tag for c in final_sect}
    # schema order puts header/footer references first inside sectPr
    for ref in reversed(refs):
        if ref.tag not in have:
            final_sect.insert(0, ref)


def para(style=None, align=None, keep=False):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass
    if align is not None:
        p.alignment = align
    if keep:
        p._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))
    return p


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def runs(p, text, bold=False, italic=False):
    """Render inline markdown (**bold**, *italic*, `code`) into runs."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"
            if p.style.font.size:
                r.font.size = Pt(p.style.font.size.pt - 1)
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            r = p.add_run(tok)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
    return p


_TRIM_DIR = Path(tempfile.mkdtemp(prefix="dvcon_fig_"))


def trimmed(path):
    """Return a copy of the figure with its white padding equalised.

    A centred paragraph centres the image *file*, not the drawing inside it, so
    baked-in asymmetric whitespace shows up as an off-centre figure on the page.
    figure4_coverage_closure.png carries 33px of left padding against 246px on
    the right (10.7% of its width), which visibly pushes the chart left and
    wastes measure. Crop to the content box, then re-pad symmetrically so the
    figure keeps a little breathing room. The source asset is never modified --
    the cropped copy goes to a temp dir.
    """
    path = Path(path)
    im = Image.open(path).convert("RGB")
    bbox = ImageChops.difference(im, Image.new("RGB", im.size, (255, 255, 255))).getbbox()
    if bbox is None:
        return str(path)
    pad = max(2, round(im.width * 0.005))
    content = im.crop(bbox)
    out = Image.new("RGB", (content.width + 2 * pad, content.height + 2 * pad),
                    (255, 255, 255))
    out.paste(content, (pad, pad))
    dest = _TRIM_DIR / path.name
    out.save(dest)
    return str(dest)


def set_widths(table, weights, num=None):
    """Fix column widths via the tblGrid; Word lays out from the grid, and
    per-cell widths alone are not enough to stop autofit from overriding.

    A weight list whose length does not match the table's column count is a
    hard error, not something to zip past: the surplus weights still count
    towards `total`, so the table silently renders narrower than the text
    measure (a stale 5-weight entry against a 4-column Table I shrank it to
    75% of the page, and a 7-weight entry against a 2-column Table IV to 34%).
    """
    ncol = len(table.columns)
    if len(weights) != ncol:
        raise SystemExit(
            f"TABLE_WEIGHTS[{num!r}] has {len(weights)} weights but the table "
            f"has {ncol} columns -- update the entry to match the markdown.")
    total = sum(weights)
    widths = [Inches(TEXT_W * w / total) for w in weights]
    table.autofit = False
    # Centre the table on the text measure. Without an explicit w:jc a table
    # defaults to left-aligned, so any rounding between the grid sum and the
    # measure hangs off the right edge instead of splitting evenly.
    tblPr = table._tbl.tblPr
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
            gc.set(qn("w:w"), str(int(w.inches * 1440)))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


lines = SRC.read_text(encoding="utf-8").split("\n")
i = 0
seen_abstract = False
pending_table = None          # caption text for the next table
in_refs = False

while i < len(lines):
    s = lines[i].strip()

    # ---------------------------------------------------------------- title
    if s.startswith("# "):
        runs(para("paper title"), s[2:])
        i += 1
        continue

    if s.startswith("*") and s.endswith("*") and not seen_abstract \
            and not s.startswith("**") and i < 8:
        runs(para("paper subtitle"), s.strip("*"))
        i += 1
        continue

    # author lines sit between the subtitle and "## Abstract"
    if re.match(r"^[A-Z][a-z]+ [A-Z][a-z]+, .*\(.*@.*\)$", s):
        runs(para("Author"), s)
        i += 1
        continue

    if not s:
        i += 1
        continue

    # ------------------------------------------------------------- headings
    if s.startswith("## "):
        title = s[3:]
        if title.lower() == "abstract":
            # The template separates the author block from the abstract with an
            # empty 'Affiliation' paragraph (6pt before, 10pt after). Reproduce it.
            para("Affiliation")
            seen_abstract = True
            i += 1
            continue
        if title.lower() in ("acknowledgment", "references"):
            runs(para("Heading 5"), title.upper())
            in_refs = title.lower() == "references"
            i += 1
            continue
        # strip the manual roman numeral -- Heading 1 auto-numbers
        runs(para("Heading 1"), re.sub(r"^[IVX]+\.\s*", "", title))
        i += 1
        continue

    if s.startswith("### "):
        # strip the manual letter -- Heading 2 auto-numbers
        runs(para("Heading 2"), re.sub(r"^[A-Z]\.\s*", "", s[4:]))
        i += 1
        continue

    # -------------------------------------------------------------- figures
    fm = re.match(r"^!\[.*?\]\((.+?)\)$", s)
    if fm:
        name = Path(fm.group(1)).name
        frac = FIG_FRAC.get(name, FIG_FRAC_DEFAULT)
        p = para(align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
        p.add_run().add_picture(trimmed(BASE / fm.group(1)),
                                width=Inches(TEXT_W * frac))
        i += 1
        continue

    # figure caption (NOT auto-numbered: keep the literal "Figure N.")
    cm = re.match(r"^\*(Figure \d+\..*)\*$", s)
    if cm:
        runs(para("figure caption", WD_ALIGN_PARAGRAPH.CENTER), cm.group(1))
        i += 1
        continue

    # --------------------------------------------------------------- tables
    tm = re.match(r"^\*\*Table ([IVX]+) — (.+?)\*\*$", s)
    if tm:
        pending_table = (tm.group(1), tm.group(2))
        i += 1
        continue

    if s.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                rows.append(cells)
            i += 1
        num, cap = pending_table if pending_table else (None, "")
        pending_table = None
        # caption ABOVE the table; 'table head' auto-numbers "TABLE I. "
        if cap:
            runs(para("table head", WD_ALIGN_PARAGRAPH.CENTER, keep=True), cap)
        ncol = max(len(r) for r in rows)
        tbl = doc.add_table(rows=0, cols=ncol)
        tbl.style = "Table Grid"
        for ri, cells in enumerate(rows):
            cells = cells + [""] * (ncol - len(cells))
            row = tbl.add_row()
            if ri == 0:
                # Repeat the header on every page a long table spills onto --
                # Table II runs to 19 rows and breaks mid-data, where an
                # unlabelled continuation is unreadable.
                row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
            centre = TABLE_CENTRE_COLS.get(num, set())
            for ci, (cell, txt) in enumerate(zip(row.cells, cells)):
                cp = cell.paragraphs[0]
                cp.style = doc.styles["table col head" if ri == 0 else "table copy"]
                if ri > 0 and ci in centre:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                runs(cp, txt)
        if num in TABLE_WEIGHTS:
            set_widths(tbl, TABLE_WEIGHTS[num], num)
        else:
            set_widths(tbl, [1.0] * ncol, num)
        # No spacer paragraph after the table: the following body paragraph
        # already carries space-before, and six spacers cost ~1/8 page against
        # the 8-page cap.
        continue

    # ----------------------------------------------------------- references
    if in_refs and re.match(r"^\[\d+\]\s", s):
        # 'references' auto-numbers "[1]" -- strip the manual marker
        runs(para("references"), re.sub(r"^\[\d+\]\s+", "", s))
        i += 1
        continue

    # ----------------------------------------------------------- body / etc
    if s.startswith("**Keywords**"):
        runs(para("key words"), "Keywords—" + s.split("—", 1)[1].strip())
        i += 1
        continue

    if s.startswith("- "):
        runs(para("bullet list"), s[2:])
        i += 1
        continue

    if seen_abstract and not any(p.style.name == "Abstract" for p in doc.paragraphs):
        p = para("Abstract")
        # Template's own run-in head: "Abstract" bold+italic, the em dash bold only.
        head = p.add_run("Abstract")
        head.bold, head.italic = True, True
        p.add_run("—").bold = True
        runs(p, s, bold=True)
        i += 1
        continue

    runs(para("Body Text", WD_ALIGN_PARAGRAPH.JUSTIFY), s)
    i += 1

doc.save(str(OUT))
print(f"wrote {OUT} (text measure {TEXT_W:.2f}in, template: {TPL.name})")
