#!/usr/bin/env python3
"""Build a camera-ready, single-column A4 DOCX from camera_ready.md.

Follows the DVCon Europe / IEEE conference template's style rules:
  - A4, SINGLE column (the DVCon Europe template is not two-column)
  - Named styles matching the template: Paper Title, Paper Subtitle, Author,
    Abstract (run-in head, bold body), Keywords, Heading 1/2/5, Caption
  - Times New Roman 10pt body
  - Table heads ABOVE tables, figure captions BELOW figures
  - Heading 5 equivalent (non-subordinate component heads) for
    ACKNOWLEDGMENT and REFERENCES
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = Path(sys.argv[1])
OUT = Path(sys.argv[2])
BASE = SRC.parent

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(10)
COL_WIDTH_IN = 4.40          # figure width; single-column template, centred
FULL_WIDTH_IN = 6.7

# Tables that must span both columns (too wide for a single column).
WIDE_TABLES = {"I", "II", "III", "IV"}

# Explicit column widths (inches), summing to FULL_WIDTH_IN. Word's autofit
# breaks long code identifiers mid-token when a column is too narrow, so the
# columns carrying identifiers are given the room they actually need.
TABLE_WIDTHS = {
    "I":   [1.30, 1.15, 1.05, 1.55, 1.65],
    "II":  [1.30, 1.55, 0.98, 0.94, 0.98, 0.95],
    "III": [0.85, 1.15, 1.15, 1.25, 1.35, 0.95],
    "IV":  [1.30, 0.90, 0.95, 0.65, 0.95, 1.00, 0.65],
}


# --------------------------------------------------------------------------
# low-level OXML helpers
# --------------------------------------------------------------------------
def set_columns(section, num):
    """Set the number of text columns on a section."""
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "180")          # ~0.125" gutter, IEEE-ish
    cols.set(qn("w:equalWidth"), "1")
    for child in list(cols):
        cols.remove(child)


def copy_page_setup(dst, src):
    dst.page_width = src.page_width
    dst.page_height = src.page_height
    dst.orientation = src.orientation
    dst.left_margin = src.left_margin
    dst.right_margin = src.right_margin
    dst.top_margin = src.top_margin
    dst.bottom_margin = src.bottom_margin


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def keep_with_next(par):
    pPr = par._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


def hanging_indent(par, left_in=0.22):
    pf = par.paragraph_format
    pf.left_indent = Inches(left_in)
    pf.first_line_indent = Inches(-left_in)


# --------------------------------------------------------------------------
# inline markdown -> runs
# --------------------------------------------------------------------------
INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_runs(par, text, size=BODY_SIZE, italic=False, bold=False):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(size.pt - 1.5)
        elif tok.startswith("*") and tok.endswith("*"):
            r = par.add_run(tok[1:-1]); r.italic = True
        else:
            r = par.add_run(tok)
        if r.font.size is None:
            r.font.size = size
        if r.font.name is None or r.font.name != "Consolas":
            r.font.name = BODY_FONT
        if italic:
            r.italic = True
        if bold:
            r.bold = True
    return par


# --------------------------------------------------------------------------
# document setup
# --------------------------------------------------------------------------
doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = BODY_SIZE
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.line_spacing = 1.0

# --------------------------------------------------------------------------
# Named paragraph styles matching the DVCon Europe template's own style names,
# so the final pour-in is "select text, pick style" rather than reformatting.
# --------------------------------------------------------------------------
from docx.enum.style import WD_STYLE_TYPE


def mkstyle(name, *, size, bold=False, italic=False, align=None,
            small_caps=False, space_before=0, space_after=0,
            first_line=None, base="Normal"):
    try:
        st = doc.styles[name]
    except KeyError:
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles[base]
    f = st.font
    f.name = BODY_FONT
    f.size = size
    # built-in Heading/Caption styles carry a theme colour; force black so the
    # document matches the template's plain black text
    f.color.rgb = RGBColor(0, 0, 0)
    f.bold = bold
    f.italic = italic
    f.small_caps = small_caps
    pf = st.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line is not None:
        pf.first_line_indent = Inches(first_line)
    return st


C = WD_ALIGN_PARAGRAPH.CENTER
J = WD_ALIGN_PARAGRAPH.JUSTIFY

mkstyle("Paper Title",    size=Pt(20), align=C, space_after=6)
mkstyle("Paper Subtitle", size=Pt(13), align=C, italic=True, space_after=10)
mkstyle("Author",         size=Pt(11), align=C, space_after=2)
mkstyle("Abstract",       size=Pt(9),  align=J, bold=True, space_after=6, first_line=0.2)
mkstyle("Keywords",       size=Pt(9),  align=J, bold=True, italic=True, space_after=9, first_line=0.2)
mkstyle("Heading 1",      size=Pt(10), align=C, bold=True, small_caps=True,
        space_before=10, space_after=4)
mkstyle("Heading 2",      size=Pt(10), italic=True, space_before=8, space_after=3)
mkstyle("Heading 5",      size=Pt(10), align=C, bold=True, small_caps=True,
        space_before=10, space_after=4)
mkstyle("Caption",        size=Pt(8),  align=C, space_before=4, space_after=3)
mkstyle("Body Text",      size=Pt(10), align=J, space_after=4, first_line=0.2)
mkstyle("Reference",      size=Pt(8),  space_after=3)

sec0 = doc.sections[0]
sec0.page_width = Inches(8.27)      # A4
sec0.page_height = Inches(11.69)
sec0.top_margin = Inches(0.75)
sec0.bottom_margin = Inches(0.75)
sec0.left_margin = Inches(0.75)
sec0.right_margin = Inches(0.75)
set_columns(sec0, 1)                # title block spans the page

lines = SRC.read_text(encoding="utf-8").split("\n")

state = {"section": sec0, "cols": 1}


def switch_columns(n):
    """Start a continuous section with n columns (no-op if already there)."""
    if state["cols"] == n:
        return
    new = doc.add_section(WD_SECTION.CONTINUOUS)
    copy_page_setup(new, state["section"])
    set_columns(new, n)
    state["section"] = new
    state["cols"] = n


def para(text="", *, align=None, size=BODY_SIZE, italic=False, bold=False,
         space_before=0, space_after=0, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_runs(p, text, size=size, italic=italic, bold=bold)
    return p


# --------------------------------------------------------------------------
# parse + emit
# --------------------------------------------------------------------------
ROMAN = re.compile(r"^(I{1,3}|IV|V|VI{0,3}|IX|X)\.\s")

i = 0
in_title_block = True
pending_table_caption = None
abstract_pending = [False]

while i < len(lines):
    raw = lines[i]
    line = raw.rstrip()
    stripped = line.strip()

    # ---- title block ------------------------------------------------------
    if in_title_block:
        if stripped.startswith("# "):
            para(stripped[2:], size=Pt(20), style="Paper Title")
            i += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*") and "Python-Based" in stripped:
            para(stripped.strip("*"), size=Pt(13), style="Paper Subtitle")
            i += 1
            continue
        # template author form: Name, Affiliation, Organization, City, Country (e-mail)
        if re.match(r"^[A-Z][a-z]+ [A-Z][a-z]+, ", stripped):
            nxt = lines[i + 1].strip() if i + 1 < len(lines) else ""
            last = not re.match(r"^[A-Z][a-z]+ [A-Z][a-z]+, ", nxt)
            p = para(stripped, size=Pt(11), style="Author")
            if last:
                p.paragraph_format.space_after = Pt(14)
            i += 1
            if last:
                in_title_block = False
            continue
        i += 1
        continue

    if not stripped:
        i += 1
        continue

    # ---- headings ---------------------------------------------------------
    if stripped.startswith("### "):
        p = para(stripped[4:], size=Pt(10), style="Heading 2")
        keep_with_next(p)
        i += 1
        continue

    if stripped.startswith("## "):
        t = stripped[3:]
        if t.lower() == "abstract":
            abstract_pending[0] = True      # run-in head, emitted with its body
            i += 1
            continue
        # numbered text heads -> Heading 1; non-subordinate component heads
        # (ACKNOWLEDGMENT, REFERENCES) -> Heading 5, per the template
        style = "Heading 1" if ROMAN.match(t) else "Heading 5"
        p = para(t.upper(), size=Pt(10), style=style)
        keep_with_next(p)
        i += 1
        continue

    # ---- figure -----------------------------------------------------------
    m = re.match(r"^!\[.*?\]\((.+?)\)$", stripped)
    if m:
        path = BASE / m.group(1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.add_run().add_picture(str(path), width=Inches(COL_WIDTH_IN))
        keep_with_next(p)
        i += 1
        continue

    # ---- figure caption (below the figure, per template) ------------------
    if re.match(r"^\*Figure \d+\.", stripped):
        p = para(stripped.strip("*"), size=Pt(8), style="Caption")
        p.paragraph_format.space_after = Pt(8)
        i += 1
        continue

    # ---- table caption (above the table, per template) --------------------
    m = re.match(r"^\*\*(Table ([IVX]+) — .+?)\*\*$", stripped)
    if m:
        pending_table_caption = (m.group(1), m.group(2))
        i += 1
        continue

    # ---- table ------------------------------------------------------------
    if stripped.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                rows.append(cells)
            i += 1

        table_num = pending_table_caption[1] if pending_table_caption else None
        wide = table_num in WIDE_TABLES

        if pending_table_caption:
            # template: table heads go ABOVE the table, styled "Caption",
            # in the form "Table I. Description"
            text = pending_table_caption[0].replace(" — ", ". ", 1)
            cap = para(text, size=Pt(8), style="Caption")
            keep_with_next(cap)
            pending_table_caption = None

        ncol = max(len(r) for r in rows)
        tbl = doc.add_table(rows=0, cols=ncol)
        tbl.style = "Table Grid"
        widths = TABLE_WIDTHS.get(table_num)
        if widths and len(widths) == ncol:
            tbl.autofit = False
            tblPr = tbl._tbl.tblPr
            layout = OxmlElement("w:tblLayout")
            layout.set(qn("w:type"), "fixed")
            tblPr.append(layout)
            tblW = OxmlElement("w:tblW")
            tblW.set(qn("w:type"), "dxa")
            tblW.set(qn("w:w"), str(int(sum(widths) * 1440)))
            tblPr.append(tblW)
            # LibreOffice/Word lay out from w:tblGrid, not from per-cell
            # widths — the grid must be rewritten or the widths are ignored.
            grid = tbl._tbl.find(qn("w:tblGrid"))
            if grid is not None:
                tbl._tbl.remove(grid)
            grid = OxmlElement("w:tblGrid")
            for w in widths:
                gc = OxmlElement("w:gridCol")
                gc.set(qn("w:w"), str(int(w * 1440)))
                grid.append(gc)
            tblPr.addnext(grid)
        else:
            widths = None
            tbl.autofit = True
        fsize = Pt(7) if wide else Pt(7.5)
        for ri, cells in enumerate(rows):
            cells = cells + [""] * (ncol - len(cells))
            wrow = tbl.add_row().cells
            for ci, ctext in enumerate(cells):
                cell = wrow[ci]
                if widths:
                    cell.width = Inches(widths[ci])
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_runs(p, ctext, size=fsize, bold=(ri == 0))
                for r in p.runs:
                    r.font.size = fsize
                if ri == 0:
                    shade(cell, "D9D9D9")

        para("", space_after=6)
        continue

    # ---- references -------------------------------------------------------
    if re.match(r"^\[\d+\]\s", stripped):
        p = para(stripped, size=Pt(8), style="Reference")
        hanging_indent(p)
        i += 1
        continue

    # ---- keywords ---------------------------------------------------------
    if stripped.startswith("**Keywords**"):
        p = doc.add_paragraph(style="Keywords")
        r = p.add_run("Keywords\u2014")
        r.bold = True; r.italic = True
        add_runs(p, stripped.split("\u2014", 1)[1].strip(), size=Pt(9), italic=True)
        i += 1
        continue

    # ---- bullets ----------------------------------------------------------
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        add_runs(p, stripped[2:], size=Pt(9.5))
        i += 1
        continue

    if re.match(r"^\d+\.\s", stripped):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        add_runs(p, re.sub(r"^\d+\.\s", "", stripped), size=Pt(9.5))
        i += 1
        continue

    # ---- ordinary paragraph ----------------------------------------------
    if abstract_pending[0]:
        # template: "Abstract" is a run-in head (bold italic) on a bold body
        p = doc.add_paragraph(style="Abstract")
        r = p.add_run("Abstract—")
        r.bold = True
        r.italic = True
        r.font.name = BODY_FONT
        r.font.size = Pt(9)
        add_runs(p, stripped, size=Pt(9), bold=True)
        abstract_pending[0] = False
        i += 1
        continue
    para(stripped, style="Body Text")
    i += 1

doc.save(str(OUT))
print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")
